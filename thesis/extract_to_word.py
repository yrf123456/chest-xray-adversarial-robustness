import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_latex(text):
    """Remove LaTeX commands and convert to plain text"""
    # Remove comments
    text = re.sub(r'%.*?\n', '\n', text)
    
    # Remove common LaTeX commands
    text = re.sub(r'\\cite\{[^}]*\}', '', text)
    text = re.sub(r'\\ref\{[^}]*\}', '', text)
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\item', '•', text)
    text = re.sub(r'\\begin\{itemize\}\[?[^\]]*\]?', '', text)
    text = re.sub(r'\\end\{itemize\}', '', text)
    text = re.sub(r'\\begin\{enumerate\}\[?[^\]]*\]?', '', text)
    text = re.sub(r'\\end\{enumerate\}', '', text)
    text = re.sub(r'\\begin\{equation\}', '', text)
    text = re.sub(r'\\end\{equation\}', '', text)
    text = re.sub(r'\\begin\{table\}.*?\\end\{table\}', '[Table]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', '[Figure]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{lstlisting\}.*?\\end\{lstlisting\}', '[Code]', text, flags=re.DOTALL)
    text = re.sub(r'\\includegraphics.*?\}', '', text)
    text = re.sub(r'\\caption\{[^}]*\}', '', text)
    text = re.sub(r'\\hspace\{[^}]*\}', '  ', text)
    text = re.sub(r'\\indent', '  ', text)
    
    # Math
    text = re.sub(r'\$([^$]*)\$', r'\1', text)
    text = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', text)
    text = re.sub(r'\\epsilon', 'ε', text)
    text = re.sub(r'\\alpha', 'α', text)
    text = re.sub(r'\\times', '×', text)
    text = re.sub(r'\\infty', '∞', text)
    text = re.sub(r'\\leq', '≤', text)
    text = re.sub(r'\\geq', '≥', text)
    text = re.sub(r'\\nabla', '∇', text)
    text = re.sub(r'\\mathcal\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\text\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\left', '', text)
    text = re.sub(r'\\right', '', text)
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    
    # Clean up
    text = re.sub(r'\{|\}', '', text)
    text = re.sub(r'~', ' ', text)
    text = re.sub(r'``|\'\'', '"', text)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    return text.strip()

def extract_sections(latex_content):
    """Extract main body sections from LaTeX"""
    sections = []
    
    # Find Introduction position first
    intro_match = re.search(r'\\section\{Introduction\}', latex_content)
    
    if not intro_match:
        return []
    
    # Search for end markers only AFTER Introduction (stop at References section)
    content_after_intro = latex_content[intro_match.start():]
    end_match = re.search(r'\\section\*?\{References\}|\\printbibliography|\\begin\{appendices\}|\\section\{Appendices\}', content_after_intro)
    
    if end_match:
        body = content_after_intro[:end_match.start()]
    else:
        body = content_after_intro
    
    # Split by sections
    section_pattern = r'\\section\{([^}]+)\}'
    subsection_pattern = r'\\subsection\{([^}]+)\}'
    subsubsection_pattern = r'\\subsubsection\{([^}]+)\}'
    
    # Find all section markers
    parts = re.split(r'(\\section\{[^}]+\}|\\subsection\{[^}]+\}|\\subsubsection\{[^}]+\})', body)
    
    current_section = None
    current_subsection = None
    
    for i, part in enumerate(parts):
        if re.match(r'\\section\{', part):
            title = re.search(r'\\section\{([^}]+)\}', part).group(1)
            sections.append(('section', title, ''))
        elif re.match(r'\\subsection\{', part):
            title = re.search(r'\\subsection\{([^}]+)\}', part).group(1)
            sections.append(('subsection', title, ''))
        elif re.match(r'\\subsubsection\{', part):
            title = re.search(r'\\subsubsection\{([^}]+)\}', part).group(1)
            sections.append(('subsubsection', title, ''))
        elif part.strip() and sections:
            # Append content to last section
            level, title, content = sections[-1]
            sections[-1] = (level, title, content + part)
    
    return sections

def create_word_doc(sections, output_path):
    """Create Word document from sections"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    section_num = 0
    subsection_num = 0
    subsubsection_num = 0
    
    for level, title, content in sections:
        if level == 'section':
            section_num += 1
            subsection_num = 0
            subsubsection_num = 0
            heading = doc.add_heading(f'{section_num}. {title}', level=1)
        elif level == 'subsection':
            subsection_num += 1
            subsubsection_num = 0
            heading = doc.add_heading(f'{section_num}.{subsection_num} {title}', level=2)
        elif level == 'subsubsection':
            subsubsection_num += 1
            heading = doc.add_heading(f'{section_num}.{subsection_num}.{subsubsection_num} {title}', level=3)
        
        # Add content
        if content.strip():
            cleaned = clean_latex(content)
            paragraphs = cleaned.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if para and para not in ['[Table]', '[Figure]', '[Code]']:
                    p = doc.add_paragraph(para)
                    p.paragraph_format.first_line_indent = Inches(0.5)
    
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")

if __name__ == '__main__':
    # Read LaTeX file
    with open('main.tex', 'r', encoding='utf-8') as f:
        latex_content = f.read()
    
    # Extract sections
    sections = extract_sections(latex_content)
    
    # Create Word document
    create_word_doc(sections, 'thesis_body.docx')
